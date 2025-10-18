#!/usr/bin/env python3
"""
Code Merger - 将指定文件夹中的代码文件整合成docx文档
支持Java/Spring Boot和Vue.js项目

改进版本，吸收了export.py的优点：
- 更好的编码处理
- 文档追加功能
- 代码清理优化
- 更健壮的错误处理
"""

import os
import sys
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import click
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError


def read_text_file_with_encoding(path: Path) -> Tuple[str, str]:
    """
    使用多种编码尝试读取文件内容
    返回 (内容, 使用的编码)
    吸收了export.py的优秀编码处理机制
    """
    encodings = ["utf-8", "gb18030", "cp936", "latin-1"]

    for enc in encodings:
        try:
            with open(path, "r", encoding=enc, errors="strict") as f:
                return f.read(), enc
        except Exception:
            continue

    # 最后的降级处理
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(), "utf-8(replace)"
    except Exception:
        return "", "unknown"


def clean_consecutive_newlines(text: str) -> str:
    """
    清理连续的空行，减少文档长度
    来自export.py的优秀优化
    """
    return re.sub(r'\n{3,}', '\n\n', text)


class CodeFile:
    """代码文件类 - 改进版本"""

    def __init__(self, file_path: Path, relative_path: str):
        self.file_path = file_path
        self.relative_path = relative_path
        self.content = ""
        self.encoding = "unknown"
        self.language = self._detect_language()
        self.read_success = False

    def _detect_language(self) -> str:
        """检测编程语言"""
        ext = self.file_path.suffix.lower()
        language_map = {
            '.java': 'Java',
            '.js': 'JavaScript',
            '.ts': 'TypeScript',
            '.vue': 'Vue',
            '.html': 'HTML',
            '.css': 'CSS',
            '.scss': 'SCSS',
            '.xml': 'XML',
            '.yml': 'YAML',
            '.yaml': 'YAML',
            '.properties': 'Properties',
            '.json': 'JSON',
            '.md': 'Markdown',
            '.py': 'Python',
            '.go': 'Go',
            '.sql': 'SQL'
        }
        return language_map.get(ext, 'Text')

    def read_content(self) -> bool:
        """读取文件内容 - 使用改进的编码处理"""
        try:
            content, encoding = read_text_file_with_encoding(self.file_path)
            if content:
                # 清理连续空行，优化文档长度
                self.content = clean_consecutive_newlines(content)
                self.encoding = encoding
                self.read_success = True
                return True
            else:
                print(f"警告: 文件内容为空或无法读取 {self.file_path} (编码: {encoding})")
                return False
        except (PermissionError, OSError) as e:
            print(f"警告: 无法读取文件 {self.file_path}: {e}")
            return False


class CodeMerger:
    """代码合并器"""

    def __init__(self):
        self.code_files: List[CodeFile] = []
        self.supported_extensions = {
            # Java/Spring Boot
            '.java', '.xml', '.yml', '.yaml', '.properties',
            # Vue.js/前端
            '.js', '.ts', '.vue', '.html', '.css', '.scss', '.json',
            # 其他常见文件
            '.md', '.py', '.go', '.sql'
        }
        self.ignore_patterns = {
            'node_modules', 'dist', 'build', 'target', '.git', '.idea',
            '.vscode', '__pycache__', '.pytest_cache', '.venv', 'venv',
            '.classpath', '.project', '.settings'
        }
        self.ignore_files = {
            '.gitignore', '.DS_Store', 'Thumbs.db', '*.log', '*.tmp'
        }

    def scan_directory(self, directory: Path) -> None:
        """扫描目录中的代码文件"""
        print(f"正在扫描目录: {directory}")

        processed = 0
        skipped = 0

        for root, dirs, files in os.walk(directory):
            # 过滤掉需要忽略的目录
            dirs[:] = [d for d in dirs if d not in self.ignore_patterns]

            for file in files:
                file_path = Path(root) / file

                # 检查文件扩展名
                if file_path.suffix.lower() not in self.supported_extensions:
                    continue

                # 检查是否需要忽略的文件
                if any(file.startswith(pattern.replace('*', '')) or file == pattern
                       for pattern in self.ignore_files):
                    continue

                # 计算相对路径
                try:
                    relative_path = str(file_path.relative_to(directory))
                except ValueError:
                    relative_path = str(file_path)

                code_file = CodeFile(file_path, relative_path)
                if code_file.read_content():
                    self.code_files.append(code_file)
                    processed += 1
                else:
                    skipped += 1

        print(f"找到 {len(self.code_files)} 个代码文件 (处理成功: {processed}, 跳过: {skipped})")

    def create_document(self, output_path: Path, project_name: str = "",
                       append_mode: bool = False) -> None:
        """创建docx文档 - 支持追加模式"""
        print(f"正在生成文档: {output_path}")

        # 文档创建逻辑，支持追加模式
        doc = self._create_or_open_document(output_path, append_mode)

        if append_mode:
            # 追加模式下添加分隔符和时间戳
            doc.add_page_break()
            doc.add_paragraph(f"追加内容时间: {self._get_current_time()}")
            doc.add_paragraph(f"追加文件数: {len(self.code_files)}")
            doc.add_paragraph("=" * 80)
        else:
            # 新文档模式 - 添加完整的标题页和目录
            self._add_title_page(doc, project_name)
            self._add_table_of_contents(doc)
            doc.add_page_break()

        # 添加代码内容
        self._add_code_contents(doc, append_mode)

        # 保存文档
        try:
            doc.save(output_path)
            print(f"文档已保存: {output_path}")
        except Exception as e:
            print(f"保存文档时出错: {e}")
            raise

    def _create_or_open_document(self, output_path: Path, append_mode: bool) -> Document:
        """创建新文档或打开现有文档（追加模式）"""
        if append_mode and output_path.exists():
            try:
                doc = Document(output_path)
                print(f"正在向现有文档追加内容: {output_path}")
                return doc
            except (PackageNotFoundError, KeyError, ValueError) as e:
                print(f"现有文档损坏或无效，将创建新文档: {e}")
                return Document()
        else:
            return Document()

    def _add_title_page(self, doc: Document, project_name: str) -> None:
        """添加标题页"""
        title = f"代码文档 - {project_name}" if project_name else "代码文档"
        doc.add_heading(title, 0)

        # 添加项目信息
        doc.add_paragraph(f"生成时间: {self._get_current_time()}")
        doc.add_paragraph(f"文件总数: {len(self.code_files)}")

        # 添加统计信息
        if self.code_files:
            languages = {}
            for code_file in self.code_files:
                lang = code_file.language
                languages[lang] = languages.get(lang, 0) + 1

            doc.add_paragraph("语言统计:")
            for lang, count in sorted(languages.items()):
                doc.add_paragraph(f"  {lang}: {count} 个文件")

    def _add_table_of_contents(self, doc: Document) -> None:
        """添加目录"""
        doc.add_heading("目录", 1)
        for i, code_file in enumerate(self.code_files, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {code_file.relative_path} ({code_file.language})")
            if code_file.encoding != "unknown":
                p.add_run(f" [编码: {code_file.encoding}]")

    def _add_code_contents(self, doc: Document, append_mode: bool) -> None:
        """添加代码内容"""
        start_index = 1

        # 如果是追加模式，需要计算正确的序号
        if append_mode:
            # 简单估算已有内容的文件数量（不完美但够用）
            existing_paragraphs = len(doc.paragraphs)
            start_index = max(1, existing_paragraphs // 10)  # 粗略估算

        for i, code_file in enumerate(self.code_files, start_index):
            # 文件标题
            doc.add_heading(f"{i}. {code_file.relative_path}", 1)

            # 文件信息
            info_para = doc.add_paragraph()
            info_para.add_run(f"语言: {code_file.language}")
            if code_file.encoding != "unknown":
                info_para.add_run(f" | 编码: {code_file.encoding}")

            # 使用改进的代码块显示（来自export.py的优化）
            self._add_code_block(doc, code_file.content)

            # 分隔符
            doc.add_paragraph("\n" + ("-" * 80) + "\n")

    def _add_code_block(self, doc: Document, content: str) -> None:
        """添加代码块 - 使用export.py的优化方法"""
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _get_current_time(self) -> str:
        """获取当前时间"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


@click.command()
@click.option('--input-dir', '-i',
              type=click.Path(exists=True, file_okay=False, dir_okay=True, path_type=Path),
              required=True,
              help='输入目录路径')
@click.option('--output', '-o',
              type=click.Path(path_type=Path),
              default='code_document.docx',
              help='输出docx文件路径 (默认: code_document.docx)')
@click.option('--project-name', '-n',
              default='',
              help='项目名称 (用于文档标题)')
@click.option('--append', '-a',
              is_flag=True,
              help='追加模式：向现有文档追加内容')
@click.option('--verbose', '-v',
              is_flag=True,
              help='显示详细输出')
def main(input_dir: Path, output: Path, project_name: str, append: bool, verbose: bool):
    """
    代码文件合并工具 - 将指定文件夹中的代码文件整合成docx文档

    支持Java/Spring Boot和Vue.js项目，自动识别文件类型并保持代码格式。
    改进版本，提供更好的编码支持和文档追加功能。
    """
    if verbose:
        print(f"输入目录: {input_dir}")
        print(f"输出文件: {output}")
        print(f"追加模式: {'是' if append else '否'}")
        if project_name:
            print(f"项目名称: {project_name}")

    try:
        merger = CodeMerger()
        merger.scan_directory(input_dir)

        if not merger.code_files:
            print("警告: 未找到任何支持的代码文件")
            return

        merger.create_document(output, project_name, append)
        print(f"✅ 文档生成完成: {output}")

    except Exception as e:
        print(f"❌ 错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()